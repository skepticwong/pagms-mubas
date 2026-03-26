try:
    from docx import Document
    import os

    # Define the path
    template_dir = os.path.join('templates', 'reports')
    file_path = os.path.join(template_dir, 'generic.docx')

    # Create directory if it doesn't exist
    if not os.path.exists(template_dir):
        os.makedirs(template_dir)

    # Create a simple document
    doc = Document()
    doc.add_heading('Generic Report Template', 0)
    doc.add_paragraph('This is a generic template for PAGMS reports.')
    doc.add_paragraph('Variables like {grant_title}, {pi_name}, etc., can be used here.')
    
    doc.save(file_path)
    print(f"Successfully created: {file_path}")

except ImportError:
    print("Error: python-docx is not installed. Please run 'pip install python-docx' and try again.")
except Exception as e:
    print(f"An error occurred: {e}")
