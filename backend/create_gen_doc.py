import os
try:
    from docx import Document
    template_dir = os.path.join('templates', 'reports')
    os.makedirs(template_dir, exist_ok=True)
    file_path = os.path.join(template_dir, 'generic.docx')
    
    doc = Document()
    doc.add_heading('Generic Report Template', 0)
    doc.add_paragraph('This is a generic template for PAGMS.')
    doc.save(file_path)
    print(f"SUCCESS: Created {file_path}")
except ImportError:
    print("FAIL: python-docx not found")
except Exception as e:
    print(f"ERROR: {e}")
